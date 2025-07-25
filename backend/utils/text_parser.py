import requests
import fitz  # PyMuPDF
from docx import Document
import logging
from pathlib import Path
from typing import Optional
import tempfile
import os
import io

logger = logging.getLogger(__name__)

class TextParser:
    """Utility class for extracting text from PDF and DOCX files."""
    
    @staticmethod
    def download_file_from_url(url: str) -> bytes:
        """
        Download file from Cloudinary URL and return file content as bytes.
        
        Args:
            url (str): Cloudinary URL of the CV file
            
        Returns:
            bytes: File content as bytes
            
        Raises:
            Exception: If download fails
        """
        try:
            response = requests.get(url, timeout=30)
            response.raise_for_status()
            return response.content
        except Exception as e:
            logger.error(f"Failed to download file from {url}: {str(e)}")
            raise Exception(f"Failed to download file: {str(e)}")
    
    @staticmethod
    def extract_text_from_pdf(file_content: bytes) -> str:
        """
        Extract text from PDF file content using PyMuPDF.
        
        Args:
            file_content (bytes): PDF file content as bytes
            
        Returns:
            str: Extracted text from all pages
        """
        try:
            # Create a temporary file to work with PyMuPDF
            with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as temp_file:
                temp_file.write(file_content)
                temp_file_path = temp_file.name
            
            try:
                # Open PDF with PyMuPDF
                doc = fitz.open(temp_file_path)
                text = ""
                
                # Extract text from all pages
                for page_num in range(doc.page_count):
                    page = doc[page_num]
                    text += page.get_text() + "\n"
                
                doc.close()
                return text.strip()
                
            finally:
                # Clean up temporary file
                os.unlink(temp_file_path)
                
        except Exception as e:
            logger.error(f"Failed to extract text from PDF: {str(e)}")
            return ""
    
    @staticmethod
    def extract_text_from_docx(file_content: bytes) -> str:
        """
        Extract text from DOCX file content using python-docx.
        
        Args:
            file_content (bytes): DOCX file content as bytes
            
        Returns:
            str: Extracted text from paragraphs and tables
        """
        try:
            # Create a temporary file to work with python-docx
            with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_file:
                temp_file.write(file_content)
                temp_file_path = temp_file.name
            
            try:
                # Open DOCX with python-docx
                doc = Document(temp_file_path)
                text = ""
                
                # Extract text from paragraphs
                for paragraph in doc.paragraphs:
                    text += paragraph.text + "\n"
                
                # Extract text from tables
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            text += cell.text + " "
                        text += "\n"
                
                return text.strip()
                
            finally:
                # Clean up temporary file
                os.unlink(temp_file_path)
                
        except Exception as e:
            logger.error(f"Failed to extract text from DOCX: {str(e)}")
            return ""
    
    @staticmethod
    def extract_text_from_cv(cv_file_path: str) -> str:
        """
        Extract text from CV file (PDF or DOCX) from Cloudinary URL.
        
        Args:
            cv_file_path (str): Cloudinary URL of the CV file
            
        Returns:
            str: Extracted text from the CV
        """
        if not cv_file_path:
            logger.warning("No CV file path provided")
            return ""
        
        try:
            # Download file content
            file_content = TextParser.download_file_from_url(cv_file_path)
            
            # Determine file type from URL or try both methods
            if cv_file_path.lower().endswith('.pdf') or 'pdf' in cv_file_path.lower():
                return TextParser.extract_text_from_pdf(file_content)
            elif cv_file_path.lower().endswith('.docx') or 'docx' in cv_file_path.lower():
                return TextParser.extract_text_from_docx(file_content)
            else:
                # Try PDF first, then DOCX
                pdf_text = TextParser.extract_text_from_pdf(file_content)
                if pdf_text:
                    return pdf_text
                
                docx_text = TextParser.extract_text_from_docx(file_content)
                if docx_text:
                    return docx_text
                
                logger.warning(f"Could not extract text from file: {cv_file_path}")
                return ""
                
        except Exception as e:
            logger.error(f"Error extracting text from CV {cv_file_path}: {str(e)}")
            return ""
    
    @staticmethod
    def extract_jd_text(job) -> str:
        """
        Extract and combine text from Job Description fields.
        
        Args:
            job: Job model instance from database
            
        Returns:
            str: Combined text from job fields
        """
        jd_text = ""
        
        if job.title:
            jd_text += f"Title: {job.title}\n"
        
        if job.description:
            jd_text += f"Description: {job.description}\n"
        
        if job.time_zone:
            jd_text += f"Time Zone: {job.time_zone}\n"
        
        if job.budget_range:
            jd_text += f"Budget Range: {job.budget_range}\n"
        
        if job.contract_duration:
            jd_text += f"Contract Duration: {job.contract_duration}\n"
        
        return jd_text.strip()